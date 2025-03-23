import streamlit as st
import ezdxf
import fitz  # PyMuPDF
import tempfile
import os
import openai
import json
import pandas as pd
from docx import Document
from io import BytesIO
from zipfile import ZipFile

st.set_page_config(page_title="Drawing Review Assistant", layout="wide")

st.title("📐 Drawing Review Assistant")

# Safe API key setup
if "OPENAI_API_KEY" in st.secrets:
    openai.api_key = st.secrets["OPENAI_API_KEY"]
else:
    st.error("❌ OpenAI API key not found. Please set it in Streamlit Cloud → Manage App → Secrets.")
    st.stop()

def extract_text_from_dxf(file_path):
    doc = ezdxf.readfile(file_path)
    msp = doc.modelspace()
    text_items = [entity.dxf.text for entity in msp if entity.dxftype() == 'TEXT']
    return "\n".join(text_items)

def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    return "\n".join([page.get_text() for page in doc])

def generate_summary_with_gpt(text):
    prompt = f"""
You are a technical assistant that reviews architectural and engineering drawing text.

Your job is to extract a structured technical summary for AEC professionals from the text below.

Please include:
- Product type and intended use
- Model number (if present)
- Manufacturer name (if present)
- Drawing title or number (if present)
- Key dimensions (organized if possible)
- Drawing views detected (Plan, Elevation, Section, Isometric, etc.)
- Scale notes or annotations (e.g. DO NOT SCALE)
- Recommended use-cases for this drawing
- Technical score (out of 10) based on clarity, completeness, and dimensional usability
- One-sentence justification for the score

Here is the raw drawing text:
\"\"\"
{text}
\"\"\"

Respond in the following JSON format:
{{
  "Product Type": "...",
  "Model Number": "...",
  "Manufacturer": "...",
  "Drawing Title/Number": "...",
  "Key Dimensions": "...",
  "Views": [...],
  "Scale Note": "...",
  "Use Cases": "...",
  "Technical Score": 0.0,
  "Score Justification": "..."
}}
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4
        )
        content = response["choices"][0]["message"]["content"]
        return json.loads(content)

    except openai.error.AuthenticationError:
        st.error("❌ OpenAI authentication failed. Check your API key.")
    except openai.error.OpenAIError as e:
        st.error(f"❌ OpenAI API error: {str(e)}")
    except json.JSONDecodeError:
        st.error("❌ GPT response could not be parsed. Check the model output format.")

    return {}

def create_word_summary(summary_dict):
    doc = Document()
    doc.add_heading("Drawing Review Summary", 0)

    # Metadata
    doc.add_heading("Drawing Metadata", level=1)
    doc.add_paragraph(f"Product Type: {summary_dict.get('Product Type', 'N/A')}")
    doc.add_paragraph(f"Model Number: {summary_dict.get('Model Number', 'N/A')}")
    doc.add_paragraph(f"Manufacturer: {summary_dict.get('Manufacturer', 'N/A')}")
    doc.add_paragraph(f"Drawing Title/Number: {summary_dict.get('Drawing Title/Number', 'N/A')}")

    # Dimensions
    doc.add_heading("Key Dimensions", level=1)
    doc.add_paragraph(summary_dict.get("Key Dimensions", "N/A"))

    # Views
    doc.add_heading("Drawing Views Detected", level=1)
    views = summary_dict.get("Views", [])
    if isinstance(views, list):
        for view in views:
            doc.add_paragraph(f"- {view}", style="List Bullet")
    else:
        doc.add_paragraph(views)

    # Scale and Annotations
    doc.add_heading("Scale and Annotations", level=1)
    doc.add_paragraph(summary_dict.get("Scale Note", "N/A"))

    # Use Cases
    doc.add_heading("Recommended Use Cases", level=1)
    doc.add_paragraph(summary_dict.get("Use Cases", "N/A"))

    # Technical Evaluation
    doc.add_heading("Technical Score", level=1)
    doc.add_paragraph(f"Score: {summary_dict.get('Technical Score', 'N/A')}/10")
    doc.add_paragraph(f"Justification: {summary_dict.get('Score Justification', 'N/A')}")

    # Final thoughts
    doc.add_paragraph("This document was generated automatically via the Drawing Review Assistant.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_excel_summary(summary_dict):
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        workbook = writer.book
        worksheet = workbook.add_worksheet("Drawing Summary")
        writer.sheets["Drawing Summary"] = worksheet

        row = 0

        # Metadata Section
        worksheet.write(row, 0, "🧾 Drawing Metadata")
        row += 1
        for label in ["Product Type", "Model Number", "Manufacturer", "Drawing Title/Number"]:
            worksheet.write(row, 0, label)
            worksheet.write(row, 1, summary_dict.get(label, ""))
            row += 1

        row += 1
        worksheet.write(row, 0, "📐 Dimensions")
        row += 1
        worksheet.write(row, 0, "Key Dimensions")
        worksheet.write(row, 1, summary_dict.get("Key Dimensions", ""))
        row += 2

        worksheet.write(row, 0, "🖼 Views Detected")
        row += 1
        views = summary_dict.get("Views", [])
        if isinstance(views, list):
            for v in views:
                worksheet.write(row, 0, "- " + v)
                row += 1
        else:
            worksheet.write(row, 0, views)
            row += 1

        row += 1
        worksheet.write(row, 0, "📏 Scale & Annotations")
        row += 1
        worksheet.write(row, 0, "Scale Note")
        worksheet.write(row, 1, summary_dict.get("Scale Note", ""))
        row += 2

        worksheet.write(row, 0, "✅ Use Cases")
        worksheet.write(row, 1, summary_dict.get("Use Cases", ""))
        row += 2

        worksheet.write(row, 0, "📊 Technical Score")
        worksheet.write(row, 1, summary_dict.get("Technical Score", ""))
        row += 1
        worksheet.write(row, 0, "Score Justification")
        worksheet.write(row, 1, summary_dict.get("Score Justification", ""))

    buffer.seek(0)
    return buffer

def create_marketing_summary(summary_dict):
    doc = Document()
    doc.add_heading("Drawing Overview", 0)
    doc.add_paragraph("This drawing showcases a high-quality, professionally detailed component suitable for architectural and engineering coordination.")
    doc.add_paragraph(f"Product Type: {summary_dict.get('Product Type', 'N/A')}")
    doc.add_paragraph(f"Key Dimensions: {summary_dict.get('Key Dimensions', 'N/A')}")
    doc.add_paragraph(f"Recommended Use: {summary_dict.get('Use Cases', 'N/A')}")
    doc.add_paragraph("This drawing meets expectations for use in submittals, BIM integration, and construction coordination workflows.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

uploaded_file = st.file_uploader("Upload a drawing file (.pdf or .dxf)", type=["pdf", "dxf"])

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_path = tmp_file.name

    if uploaded_file.name.endswith(".dxf"):
        raw_text = extract_text_from_dxf(tmp_path)
    else:
        raw_text = extract_text_from_pdf(tmp_path)

    st.subheader("📄 Extracted Text Preview")
    st.text_area("Raw Drawing Text", raw_text[:3000], height=300)

    st.subheader("📊 GPT-Powered Technical Summary")
    summary = generate_summary_with_gpt(raw_text)

    if isinstance(summary, dict) and summary:
        for key, value in summary.items():
            st.write(f"**{key}:** {value}")

        word_buffer = create_word_summary(summary)
        excel_buffer = create_excel_summary(summary)
        marketing_buffer = create_marketing_summary(summary)

        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, "w") as zip_file:
            zip_file.writestr("technical_summary.docx", word_buffer.getvalue())
            zip_file.writestr("drawing_summary.xlsx", excel_buffer.getvalue())
            zip_file.writestr("marketing_summary.docx", marketing_buffer.getvalue())
        zip_buffer.seek(0)

        st.download_button(
            "📦 Download All Files (ZIP)",
            zip_buffer,
            file_name="drawing_review_outputs.zip",
            mime="application/zip"
        )
    else:
        st.error("⚠️ No valid summary was generated. Please check the GPT response.")
